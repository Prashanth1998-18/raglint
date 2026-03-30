"""Document and chunk export parsers."""

from __future__ import annotations

import csv
import io
import json
from datetime import datetime
from pathlib import Path
from typing import Any

import fitz
from docx import Document as DocxDocument
from fastapi import UploadFile

from app.errors import (
    ChunkImportError,
    ChunkLimitError,
    DocumentParsingError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from app.models.document import Chunk, ChunkPosition, Document
from app.services.metadata import (
    build_document_metadata,
    parse_markdown_frontmatter,
    sanitize_embedded_metadata,
)


class DocumentParser:
    """Parse supported documents into a normalized internal model."""

    supported_extensions = {".pdf", ".docx", ".md", ".markdown", ".txt"}

    async def parse_upload(
        self,
        upload_file: UploadFile,
        *,
        client_modified_at: datetime | None = None,
    ) -> Document:
        """Parse a FastAPI upload into a Document model."""
        filename = upload_file.filename or "uploaded-file"
        extension = Path(filename).suffix.lower()
        if extension not in self.supported_extensions:
            raise UnsupportedFileTypeError(
                f"File {filename} is not supported. RAGLint accepts PDF, DOCX, Markdown, and TXT files."
            )

        raw_bytes = await upload_file.read()
        if not raw_bytes:
            raise EmptyDocumentError(f"File {filename} produced no extractable text and was excluded from analysis.")

        text = ""
        embedded_metadata: dict[str, object] = {}
        frontmatter: dict[str, object] = {}

        try:
            if extension == ".pdf":
                text, embedded_metadata = self._parse_pdf(raw_bytes)
            elif extension == ".docx":
                text, embedded_metadata = self._parse_docx(raw_bytes)
            elif extension in {".md", ".markdown"}:
                markdown_text = raw_bytes.decode("utf-8", errors="replace").replace("\r\n", "\n")
                frontmatter, text = parse_markdown_frontmatter(markdown_text)
            else:
                text = raw_bytes.decode("utf-8", errors="replace")
        except Exception as exc:  # pragma: no cover - defensive library wrapper
            raise DocumentParsingError("unable to extract text") from exc

        embedded_created_at = self._extract_datetime(embedded_metadata, "created", "creationDate")
        embedded_modified_at = self._extract_datetime(embedded_metadata, "modified", "modDate")

        normalized_text = text.strip()
        if not normalized_text:
            raise EmptyDocumentError(f"File {filename} produced no extractable text and was excluded from analysis.")

        metadata = build_document_metadata(
            filename=filename,
            file_size=len(raw_bytes),
            created_at=embedded_created_at,
            modified_at=client_modified_at or embedded_modified_at,
            embedded=embedded_metadata,
            frontmatter=frontmatter,
        )

        return Document(
            text=normalized_text,
            metadata=metadata,
            source_path=filename,
        )

    def _parse_pdf(self, raw_bytes: bytes) -> tuple[str, dict[str, object]]:
        """Extract text and document properties from a PDF."""
        with fitz.open(stream=raw_bytes, filetype="pdf") as pdf_document:
            text = "\n\n".join(page.get_text("text") for page in pdf_document)
            metadata = sanitize_embedded_metadata(pdf_document.metadata or {})
        return text, metadata

    def _parse_docx(self, raw_bytes: bytes) -> tuple[str, dict[str, object]]:
        """Extract text and core properties from a DOCX file."""
        buffer = io.BytesIO(raw_bytes)
        document = DocxDocument(buffer)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

        core_properties = document.core_properties
        metadata = sanitize_embedded_metadata(
            {
                "author": core_properties.author,
                "category": core_properties.category,
                "comments": core_properties.comments,
                "content_status": core_properties.content_status,
                "created": core_properties.created,
                "identifier": core_properties.identifier,
                "keywords": core_properties.keywords,
                "language": core_properties.language,
                "last_modified_by": core_properties.last_modified_by,
                "last_printed": core_properties.last_printed,
                "modified": core_properties.modified,
                "revision": core_properties.revision,
                "subject": core_properties.subject,
                "title": core_properties.title,
                "version": core_properties.version,
            }
        )

        return text, metadata

    def _extract_datetime(
        self,
        metadata: dict[str, object],
        *keys: str,
    ) -> datetime | None:
        """Extract common ISO or PDF datetime formats from embedded metadata."""
        for key in keys:
            value = metadata.get(key)
            if not isinstance(value, str):
                continue
            parsed = self._parse_datetime_string(value)
            if parsed is not None:
                return parsed
        return None

    def _parse_datetime_string(self, value: str) -> datetime | None:
        """Parse ISO 8601 or PDF metadata datetime strings."""
        try:
            return datetime.fromisoformat(value.replace("Z", "+00:00"))
        except ValueError:
            pass

        pdf_value = value.strip()
        if not pdf_value.startswith("D:"):
            return None

        digits = "".join(character for character in pdf_value[2:] if character.isdigit())
        for pattern_length in (14, 12, 8):
            if len(digits) < pattern_length:
                continue
            try:
                if pattern_length == 14:
                    return datetime.strptime(digits[:14], "%Y%m%d%H%M%S")
                if pattern_length == 12:
                    return datetime.strptime(digits[:12], "%Y%m%d%H%M")
                return datetime.strptime(digits[:8], "%Y%m%d")
            except ValueError:
                continue
        return None


class ChunkExportParser:
    """Normalize JSON or CSV chunk exports into the internal chunk model."""

    supported_extensions = {".json", ".csv"}

    def __init__(self, *, max_chunks: int = 5000) -> None:
        self.max_chunks = max_chunks

    async def parse_upload(self, upload_file: UploadFile) -> list[Chunk]:
        """Parse and normalize an uploaded chunk export."""
        filename = upload_file.filename or "chunks-export"
        extension = Path(filename).suffix.lower()
        if extension not in self.supported_extensions:
            raise ChunkImportError(
                "The chunks export file could not be read. Please ensure it is valid JSON or CSV."
            )

        raw_bytes = await upload_file.read()
        if not raw_bytes:
            raise ChunkImportError("The chunks export file could not be read. Please ensure it is valid JSON or CSV.")

        if extension == ".json":
            rows = self._load_json(raw_bytes)
        else:
            rows = self._load_csv(raw_bytes)

        normalized_chunks: list[Chunk] = []
        for index, row in enumerate(rows):
            chunk = self._normalize_chunk_row(row, index)
            if chunk is not None:
                normalized_chunks.append(chunk)

        if rows and not normalized_chunks:
            raise ChunkImportError("Chunks export is missing required 'text' field.")
        if len(normalized_chunks) > self.max_chunks:
            raise ChunkLimitError(
                f"Your chunks export contains {len(normalized_chunks)} chunks. The maximum is 5,000. Please reduce the export size."
            )
        return normalized_chunks

    def _load_json(self, raw_bytes: bytes) -> list[dict[str, Any]]:
        """Load rows from common JSON chunk export shapes."""
        try:
            payload = json.loads(raw_bytes.decode("utf-8", errors="replace"))
        except json.JSONDecodeError as exc:
            raise ChunkImportError(
                "The chunks export file could not be read. Please ensure it is valid JSON or CSV."
            ) from exc
        if isinstance(payload, list):
            return [item for item in payload if isinstance(item, dict)]
        if isinstance(payload, dict):
            for key in ("chunks", "data", "items", "results"):
                value = payload.get(key)
                if isinstance(value, list):
                    return [item for item in value if isinstance(item, dict)]
        raise ChunkImportError("The chunks export file could not be read. Please ensure it is valid JSON or CSV.")

    def _load_csv(self, raw_bytes: bytes) -> list[dict[str, Any]]:
        """Load rows from a CSV chunk export."""
        decoded = raw_bytes.decode("utf-8-sig", errors="replace")
        reader = csv.DictReader(io.StringIO(decoded))
        if not reader.fieldnames:
            raise ChunkImportError("The chunks export file could not be read. Please ensure it is valid JSON or CSV.")
        return [dict(row) for row in reader]

    def _normalize_chunk_row(self, row: dict[str, Any], index: int) -> Chunk | None:
        """Map a source row into the shared chunk format."""
        text = self._first_string(
            row,
            "text",
            "content",
            "chunk_text",
            "page_content",
            "body",
        )
        if not text:
            return None

        parent_document_id = self._first_string(
            row,
            "parent_document_id",
            "document_id",
            "doc_id",
            "source_id",
        )
        parent_document_name = self._first_string(
            row,
            "parent_document_name",
            "document_name",
            "filename",
            "source",
            "source_path",
            "document",
        )
        chunk_index = self._first_int(row, "chunk_index", "index", "position", default=index)
        start_char = self._first_int(row, "start_char", "start", default=0)
        end_char = self._first_int(row, "end_char", "end", default=start_char + len(text))
        metadata, embedding = self._extract_chunk_metadata(row)

        return Chunk(
            text=text,
            parent_document_id=parent_document_id,
            parent_document_name=parent_document_name,
            position=ChunkPosition(
                chunk_index=chunk_index,
                start_char=max(0, start_char),
                end_char=max(start_char, end_char),
            ),
            metadata=metadata,
            embedding=embedding,
        )

    def _extract_chunk_metadata(self, row: dict[str, Any]) -> tuple[dict[str, object], list[float] | None]:
        """Collect extra row fields, metadata, and any provided embedding vector."""
        metadata: dict[str, object] = {}
        embedding = self._coerce_embedding(
            row.get("embedding") or row.get("embeddings") or row.get("vector")
        )
        reserved_keys = {
            "text",
            "content",
            "chunk_text",
            "page_content",
            "body",
            "parent_document_id",
            "document_id",
            "doc_id",
            "source_id",
            "parent_document_name",
            "document_name",
            "filename",
            "source",
            "source_path",
            "document",
            "chunk_index",
            "index",
            "position",
            "start_char",
            "start",
            "end_char",
            "end",
            "metadata",
            "embedding",
            "embeddings",
            "vector",
        }

        for key, value in row.items():
            if key in reserved_keys or value in (None, ""):
                continue
            metadata[key] = value

        metadata_value = row.get("metadata")
        if isinstance(metadata_value, dict):
            metadata.update(metadata_value)
        elif isinstance(metadata_value, str):
            try:
                decoded = json.loads(metadata_value)
            except json.JSONDecodeError:
                decoded = None
            if isinstance(decoded, dict):
                metadata.update(decoded)

        if embedding is None:
            embedding = self._coerce_embedding(
                metadata.pop("embedding", None)
                or metadata.pop("embeddings", None)
                or metadata.pop("vector", None)
            )
        else:
            metadata.pop("embedding", None)
            metadata.pop("embeddings", None)
            metadata.pop("vector", None)

        return metadata, embedding

    def _first_string(self, row: dict[str, Any], *keys: str) -> str | None:
        """Return the first non-empty string-like value for the provided keys."""
        for key in keys:
            value = row.get(key)
            if value is None:
                continue
            if isinstance(value, str):
                cleaned = value.strip()
                if cleaned:
                    return cleaned
                continue
            cleaned = str(value).strip()
            if cleaned:
                return cleaned
        return None

    def _first_int(self, row: dict[str, Any], *keys: str, default: int) -> int:
        """Return the first parseable integer value for the provided keys."""
        for key in keys:
            value = row.get(key)
            if value in (None, ""):
                continue
            try:
                return int(value)
            except (TypeError, ValueError):
                continue
        return default

    def _coerce_embedding(self, value: object) -> list[float] | None:
        """Normalize a stored embedding from JSON or CSV formats."""
        if value in (None, ""):
            return None

        if isinstance(value, list):
            normalized: list[float] = []
            for item in value:
                try:
                    normalized.append(float(item))
                except (TypeError, ValueError):
                    return None
            return normalized or None

        if isinstance(value, str):
            try:
                decoded = json.loads(value)
            except json.JSONDecodeError:
                stripped = value.strip().strip("[]")
                if not stripped:
                    return None
                try:
                    return [float(item.strip()) for item in stripped.split(",") if item.strip()]
                except ValueError:
                    return None
            return self._coerce_embedding(decoded)

        return None
